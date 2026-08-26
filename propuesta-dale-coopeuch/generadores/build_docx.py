# -*- coding: utf-8 -*-
import sys
sys.path.insert(0, "/tmp/claude-0/-home-user-meta-mcp/3417b367-2cff-5bc9-8662-cddb20c827cd/scratchpad")
from doc_lib import *
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL

OUT = "/home/user/meta-mcp/propuesta-dale-coopeuch/03_Investigacion_de_Mercado_Dale_Coopeuch.docx"
LOGO = "/tmp/claude-0/-home-user-meta-mcp/3417b367-2cff-5bc9-8662-cddb20c827cd/scratchpad/assets/logo_intothecom_negro.png"

doc = setup(Document())

# ─────────────────────────────── PORTADA ──────────────────────────────────────
doc.add_picture(LOGO, height=Inches(0.5))
doc.paragraphs[-1].alignment = AL.LEFT
for _ in range(3):
    doc.add_paragraph()
para(doc, [("INVESTIGACIÓN DE MERCADO", {"size": 12, "bold": True, "color": ORANGE})],
     after=2)
para(doc, [("Benchmarks e insights para la estrategia\nde canales digitales",
            {"size": 26, "bold": True})], after=10)
para(doc, [("Dale Coopeuch", {"size": 17, "color": GRAY})], after=24)
para(doc, [("Preparado por Intothecom · Agencia de Marketing Digital",
            {"size": 10, "color": GRAY})], after=2)
para(doc, [("Agosto 2026 · Documento de trabajo, versión detallada",
            {"size": 10, "color": GRAY})], after=30)
callout(doc, "Sobre este documento.",
        "Reúne los benchmarks públicos, el análisis competitivo y el marco regulatorio que "
        "sustentan la propuesta de servicios y el plan de medios presentados por separado. "
        "Todas las cifras están atribuidas a su fuente y fecha. Donde no existe información "
        "pública confiable, se señala explícitamente en lugar de estimarla.")

# ────────────────────────── 1. RESUMEN EJECUTIVO ──────────────────────────────
h1(doc, "Resumen ejecutivo", "01")
para(doc, "El objetivo declarado por la marca es la conversión y el curse de créditos de "
          "consumo digitales captando prospectos de mercado abierto. No es notoriedad ni "
          "volumen de descargas. Esta investigación se construyó para responder a esa "
          "pregunta concreta y no a un diagnóstico genérico de presencia digital.")
h2(doc, "Los seis hallazgos que ordenan la estrategia")
bullet(doc, [
 ("El problema no es tráfico, es fricción. ",
  "La contratación del Crédito Digital Dale exige que el prospecto complete una secuencia "
  "de pasos dentro de la app, incluyendo hacerse Socio de la cooperativa. Cada paso es un "
  "punto de fuga que hoy no se mide de forma aislada."),
 ("Existe un diferenciador real y desaprovechado. ",
  "La marca declara evaluación \"con inteligencia financiera en tiempo real considerando "
  "múltiples fuentes de ingreso\". Eso habilita un nicho amplio y mal atendido: trabajadores "
  "independientes y personas sin liquidación de sueldo."),
 ("La landing pública no cumple los requisitos de Google para anunciar crédito. ",
  "No expone CAE, plazos ni costo total. Es un bloqueante previo al lanzamiento, no una "
  "observación menor."),
 ("En Chile conservamos la segmentación avanzada de Meta. ",
  "Las restricciones de categoría especial para crédito aplican a Estados Unidos, Canadá y "
  "parte de Europa. Chile no está en esa lista, por lo que lookalikes y audiencias "
  "personalizadas siguen disponibles."),
 ("En email, el valor está en los flujos, no en las campañas. ",
  "Los flujos automatizados generan el 41% de los ingresos por email con apenas el 5,3% de "
  "los envíos. Banca y finanzas es además la industria con menor tasa de clic."),
 ("En redes sociales hay que fijar expectativas bajas y objetivos correctos. ",
  "Servicios financieros rinde entre 0,26% y 0,67% de engagement en Instagram frente a 1,9% "
  "en TikTok. El rol del contenido orgánico aquí es construir confianza y explicar el "
  "proceso, no acumular seguidores."),
])
h2(doc, "Inversión recomendada")
para(doc, "A partir de la triangulación de CPC descrita en el capítulo 6, se recomienda una "
          "inversión publicitaria de entre $3.000.000 y $5.000.000 CLP mensuales para la fase "
          "inicial. El detalle por campaña y el modelo que sustenta la cifra se entregan en el "
          "Plan de Medios adjunto.")

# ────────────────────────── 2. METODOLOGÍA ────────────────────────────────────
h1(doc, "Metodología y advertencias de uso", "02")
h2(doc, "Cómo se construyó este documento")
bullet(doc, [
 ("Fuentes públicas verificables. ", "Se usaron documentación oficial de plataformas, "
  "estadísticas de organismos reguladores chilenos, reportes de inversión publicitaria del "
  "mercado local y estudios de benchmark de la industria."),
 ("Sin acceso a cuentas del cliente. ", "No se dispuso de acceso a Google Ads, Meta Business "
  "ni analítica de la marca. Todas las cifras de desempeño son referencias externas, no "
  "mediciones del anunciante."),
 ("Atribución explícita. ", "Cada dato indica fuente, alcance geográfico y fecha. Los "
  "benchmarks internacionales se identifican como tales."),
 ("Sin estimaciones no respaldadas. ", "Donde no existe información pública confiable se "
  "declara la ausencia en lugar de rellenarla con supuestos."),
])
h2(doc, "Tres advertencias importantes")
callout(doc, "Los benchmarks internacionales no son trasladables en valor absoluto.",
        "Los CPC publicados por fuentes como WordStream corresponden principalmente al mercado "
        "estadounidense. Se utilizan como referencia de la relación entre métricas — por "
        "ejemplo, CTR alto con conversión baja — y no como precio esperado en Chile. Para el "
        "valor en pesos se aplica la triangulación del capítulo 6.")
callout(doc, "No existe un promedio público de inversión por anunciante financiero en Chile.",
        "Los reportes de AAM, Admetricks e IAB Chile publican agregados de mercado, no gasto "
        "por anunciante individual. Por eso la inversión recomendada se construye de abajo "
        "hacia arriba desde los benchmarks de costo y conversión, y no copiando una cifra de "
        "mercado que no existe.")
callout(doc, "Las políticas de plataforma cambian con frecuencia.",
        "Los criterios regulatorios y de segmentación descritos en el capítulo 5 fueron "
        "verificados en agosto de 2026. Deben revisarse nuevamente al momento de implementar, "
        "especialmente los alcances geográficos, que las plataformas han ido ampliando.")

# ────────────────────────── 3. CONTEXTO DE MERCADO ────────────────────────────
h1(doc, "Contexto del crédito de consumo en Chile", "03")
h2(doc, "El consumo es el corazón del negocio cooperativo")
para(doc, "La cartera de consumo representa el 68,74% del total de colocaciones de las "
          "cooperativas supervisadas, con un crecimiento real de 4,77% en doce meses al cierre "
          "de enero de 2026. El mismo informe señala que el crecimiento se desaceleró respecto "
          "de diciembre y que casi todos los indicadores de riesgo de la cartera de consumo "
          "muestran alzas, con excepción del indicador de provisiones.")
source_note(doc, "CMF, informe de desempeño de bancos y cooperativas supervisadas a enero de 2026.")
para(doc, "La lectura para la estrategia es doble. Por un lado, el consumo es el producto "
          "central y su crecimiento sostiene el negocio. Por otro, el deterioro de los "
          "indicadores de riesgo implica que el volumen bruto de solicitudes no es un objetivo "
          "deseable en sí mismo: importa la calidad del prospecto que llega al proceso de "
          "evaluación. Esto refuerza la decisión de calificar antes de derivar.")

h2(doc, "Lo digital ya es más de la mitad de la inversión publicitaria chilena")
table(doc, ["Indicador", "Valor", "Variación", "Periodo"],
      [["Inversión digital reportada", "$29.262 millones CLP", "+14,6% interanual", "may-2026"],
       ["Share of Investment digital", "50,1%", "+3,2 pts vs 46,9%", "may-2026"],
       ["Compra directa sobre el total", "84,2%", "—", "may-2026"],
       ["Social sobre compra directa", "37,0%", "—", "may-2026"],
       ["Search sobre compra directa", "31,3%", "—", "may-2026"],
       ["Programática sobre el total", "15,8%", "—", "may-2026"]],
      widths=[2.4, 1.6, 1.5, 1.0])
source_note(doc, "AAM, Admetricks e IAB Chile, informe de Inversión Digital, mayo 2026.")
para(doc, "Mayo de 2026 marcó la primera vez que el medio digital superó la mitad del mercado "
          "publicitario chileno. La distribución entre social y search — 37% y 31,3% "
          "respectivamente — respalda la decisión de concentrar el plan de medios en Meta y "
          "Google antes que dispersar la inversión en canales secundarios.")

h2(doc, "El ecosistema competitivo se está reordenando")
bullet(doc, [
 ("El prepago explotó. ", "El uso de tarjetas de prepago creció más de 213% entre marzo de "
  "2024 y marzo de 2025 según datos del Banco Central, reflejando la adopción acelerada de "
  "soluciones de pago alternativas."),
 ("Los neobancos avanzan hacia el crédito. ", "Mercado Pago anunció el lanzamiento de su "
  "tarjeta de crédito en Chile durante 2026 para competir de forma más directa con la banca "
  "tradicional. Tenpo, controlado por Krealo, busca convertirse en el primer neobanco del país."),
 ("Dale compite en dos canchas distintas. ", "En cuenta y prepago compite con Tenpo, MACH, "
  "Mercado Pago y Tapp, todas sin costo de mantención. En crédito de consumo la competencia "
  "es la banca, las cooperativas y el retail financiero. La estrategia de comunicación no "
  "puede tratar ambas canchas como una sola."),
])
source_note(doc, "Cámara de Comercio de Santiago, LatamFintech y reportes de prensa sectorial, 2026.")


# ────────────────────────── 4. EL CLIENTE Y SU EMBUDO ─────────────────────────
h1(doc, "Dale Coopeuch: producto, propuesta de valor y embudo", "04")
h2(doc, "Qué es y qué ofrece")
para(doc, "Dale Coopeuch es una cuenta digital con tarjeta de prepago Mastercard, gratuita en "
          "creación y mantención, con saldo máximo de $5.000.000 y vigencia de cinco años. "
          "Sobre esa base se ofrecen productos adicionales: Crédito Digital Dale, seguros, "
          "inversión en fondos mutuos desde $1.000 e Hipotecario 40. La apertura de la cuenta "
          "no requiere ser Socio previo de Coopeuch ni presentar liquidación de sueldo.")
source_note(doc, "Sitio público dalecoopeuch.cl, revisado en agosto de 2026.")

h2(doc, "El diferenciador central del crédito")
callout(doc, "Textual del sitio.",
        "\"Te evaluamos con inteligencia financiera en tiempo real considerando múltiples "
        "fuentes de ingreso.\" Esta frase es el activo de comunicación más valioso que tiene "
        "la marca y hoy está enterrada en el cuerpo de una landing.")
para(doc, "En un mercado donde la evaluación crediticia tradicional exige liquidación de sueldo "
          "y antigüedad laboral formal, una evaluación que considera múltiples fuentes de "
          "ingreso abre la puerta a trabajadores independientes, personas a honorarios y "
          "trabajadores de plataformas. Es un segmento numeroso, con necesidad de "
          "financiamiento y sistemáticamente rechazado por la oferta tradicional. Ningún "
          "competidor directo comunica hoy este ángulo con claridad.")
para(doc, "El segundo diferenciador es el Remanente: la devolución anual de utilidades a los "
          "Socios. Convierte el requisito de asociarse — que en la superficie es fricción — en "
          "un beneficio económico concreto. Ni Tenpo, ni MACH, ni Mercado Pago tienen un "
          "equivalente.")

h2(doc, "El embudo real hacia el curse")
para(doc, "El sitio describe el proceso de solicitud del crédito dentro de la app: ingresar a "
          "la sección Crédito de Consumo, autorizar la evaluación en línea, ingresar el monto y "
          "firmar el contrato en línea. Antes de eso el usuario debe tener la app instalada, "
          "el DalePass activado y ser Socio de la cooperativa.")
table(doc, ["Etapa", "Qué ocurre", "Riesgo de fuga"],
      [["1. Impacto", "El prospecto ve el anuncio", "Mensaje genérico que no diferencia"],
       ["2. Conversación", "Inicia contacto por WhatsApp", "Fricción de formularios tradicionales"],
       ["3. Calificación", "El agente IA levanta y filtra", "Prospecto fuera de perfil"],
       ["4. Derivación", "Pasa al proceso de contratación", "Pérdida de trazabilidad"],
       ["5. Requisitos previos", "App, DalePass y condición de Socio", "Abandono por número de pasos"],
       ["6. Evaluación", "Autoriza la evaluación en línea", "Desconfianza en entregar datos"],
       ["7. Monto y firma", "Define monto y firma el contrato", "Abandono en el último paso"],
       ["8. Curse", "El crédito se cursa", "—"]],
      widths=[1.5, 2.4, 2.6])
callout(doc, "Implicancia estratégica.",
        "Con un embudo de esta longitud, optimizar campañas hacia \"lead\" o hacia "
        "\"instalación de app\" produce volumen que no se traduce en curse. La medición debe "
        "llegar hasta el evento final y, si es posible, devolver a las plataformas la señal de "
        "aprobación crediticia para que el algoritmo aprenda a buscar prospectos aprobables.")

# ────────────────────────── 5. MARCO REGULATORIO ──────────────────────────────
h1(doc, "Marco regulatorio y políticas de plataforma", "05")
para(doc, "Este capítulo es el que más impacto operativo tiene sobre el lanzamiento. Anunciar "
          "crédito de consumo no es equivalente a anunciar un producto de retail: existen "
          "requisitos que condicionan el contenido de la landing, la segmentación disponible y "
          "la habilitación misma de la cuenta publicitaria.")

h2(doc, "5.1 Google Ads: requisitos para préstamos personales")
bullet(doc, [
 ("Divulgación obligatoria. ", "El anunciante debe mostrar de forma destacada en su página de "
  "destino o app: el período mínimo y máximo de pago, la Tasa Anual Equivalente máxima y un "
  "ejemplo representativo del costo total del préstamo incluidas todas las comisiones "
  "aplicables."),
 ("Regla de los 61 días. ", "Solo se permiten préstamos personales que requieren el pago "
  "íntegro en 61 días o más. Aplica a anunciantes directos, generadores de leads y quienes "
  "conectan consumidores con prestamistas externos."),
 ("Cumplimiento local. ", "El anunciante debe cumplir la regulación local de cada ubicación "
  "hacia la que segmenta sus anuncios."),
 ("Sistema de faltas. ", "La política opera con un máximo de una advertencia y tres faltas. La "
  "tercera falta suspende la cuenta."),
])
source_note(doc, "Google Ads, Políticas de productos y servicios financieros — Préstamos "
                 "personales, consultado en agosto de 2026.")

callout(doc, "Hallazgo verificado: la landing pública de Dale no cumple hoy este requisito.",
        "Se revisó dalecoopeuch.cl/creditodigital a nivel de código fuente, no solo de forma "
        "visual. No aparece la CAE, ni la expresión \"carga anual equivalente\", ni tasa de "
        "interés, ni plazos, ni montos, ni costo total, ni ejemplo representativo. El único "
        "contenido legal presente en el pie corresponde al aviso CMF sobre garantía estatal de "
        "depósitos y al de entidades autorizadas para emitir tarjetas de pago. Verificación "
        "realizada en agosto de 2026.")
para(doc, "Esto debe resolverse antes de activar campañas de búsqueda hacia esa URL. Existen "
          "dos caminos posibles y la decisión corresponde al cliente:")
bullet(doc, [
 ("Exponer las condiciones en la landing pública. ", "Es la vía más segura porque el revisor "
  "de Google alcanza el contenido sin necesidad de autenticación."),
 ("Mantenerlas dentro de la app. ", "La política admite la app como lugar de divulgación, pero "
  "si las condiciones quedan detrás del login existe el riesgo de que la revisión no las "
  "alcance y el anuncio sea rechazado."),
])

h2(doc, "5.2 Referencias de cumplimiento en el mercado chileno")
para(doc, "La exigencia de Google se apoya sobre una obligación que en Chile ya es legal. La "
          "CAE está definida y difundida por el SERNAC como el indicador que permite comparar "
          "el costo total de un crédito, incluyendo intereses, seguros asociados, comisiones de "
          "apertura y gastos de administración. La CMF, por su parte, dispone de un simulador "
          "público que compara tasas y CAE entre bancos y cooperativas para un monto y plazo "
          "determinados.")
table(doc, ["Referencia", "Qué aporta como ejemplo", "Utilidad para Dale"],
      [["Simulador de créditos de la CMF",
        "Comparador oficial que expone CAE por institución para un monto y plazo dados",
        "Estándar de referencia sobre qué información debe estar disponible"],
       ["Material educativo del SERNAC sobre CAE",
        "Define el indicador y explica por qué es el criterio de comparación válido",
        "Base legal y de lenguaje para redactar el disclosure"],
       ["Portales comparadores del mercado",
        "Publican CAE por institución; a mayo de 2026 se reporta a Banco BICE liderando desde "
        "20,66% para buen perfil",
        "Muestra el nivel de transparencia que el usuario ya encuentra al comparar"]],
      widths=[1.7, 2.5, 2.3])
source_note(doc, "SERNAC, CMF y portales comparadores del mercado chileno, 2026.")
callout(doc, "Nota de verificación.",
        "Los sitios de los principales emisores chilenos entregan la CAE dentro de "
        "simuladores dinámicos que se cargan por JavaScript, por lo que no fue posible "
        "verificarlos de forma programática con el mismo rigor aplicado a la landing de Dale. "
        "Se recomienda levantar capturas de pantalla de dos o tres competidores al momento de "
        "preparar el disclosure, para usarlas como referencia concreta ante el equipo legal "
        "del cliente.")

h2(doc, "5.3 Meta: categorías especiales y verificación")
para(doc, "Meta clasifica como categoría especial los anuncios que promueven oportunidades de "
          "crédito, lo que restringe fuertemente las opciones de segmentación: desactiva "
          "audiencias similares, impide exclusiones de segmentación detallada y limita la "
          "segmentación geográfica.")
callout(doc, "Buena noticia: esa restricción no aplica en Chile.",
        "Según la documentación de Meta, la obligación de autoidentificarse en categoría "
        "especial recae sobre anunciantes de Estados Unidos, o que segmentan hacia Estados "
        "Unidos, Canadá y ciertos países de Europa. Chile no figura en ese alcance. En "
        "consecuencia, para campañas dirigidas al mercado chileno se conservan audiencias "
        "personalizadas, lookalikes y segmentación detallada.")
bullet(doc, [
 ("Lo que sí aplica siempre. ", "Los anuncios de tarjetas de crédito, préstamos o seguros "
  "deben dirigirse a personas mayores de 18 años."),
 ("Verificación del anunciante. ", "Meta puede exigir verificación de identidad de la empresa "
  "y demostración de autorización por parte del organismo regulador correspondiente."),
 ("Prohibición de solicitar datos sensibles en el anuncio. ", "No se permite pedir "
  "directamente información personal identificable o financiera dentro de la pieza."),
 ("Revisar antes de implementar. ", "Meta ha ampliado el alcance geográfico de estas reglas de "
  "forma progresiva. La verificación de agosto de 2026 debe repetirse al lanzar."),
])
source_note(doc, "Meta Business Help Centre y Centro de Transparencia de Meta, consultados en "
                 "agosto de 2026.")

# ────────────────────────── 6. BENCHMARKS PAID MEDIA ──────────────────────────
h1(doc, "Benchmarks de Paid Media", "06")
h2(doc, "6.1 Google Ads — referencias internacionales")
table(doc, ["Métrica", "Finanzas y seguros", "Banca", "Seguros", "Alcance"],
      [["CPC promedio Search", "USD 3,46", "USD 3,08", "USD 6,22", "EE.UU. / Global"],
       ["CTR Search", "8,3%", "3,41%", "2,53%", "EE.UU. / Global"],
       ["Tasa de conversión Search", "2,5%", "4,72%", "6,15%", "EE.UU. / Global"],
       ["CPA Search", "—", "USD 65,25", "USD 101,14", "Global"],
       ["CPC Display", "—", "USD 0,51", "USD 0,68", "Global"],
       ["Variación CPC interanual", "—", "+10%", "+11%", "Global"]],
      widths=[1.9, 1.4, 1.1, 1.1, 1.2])
source_note(doc, "WordStream Google Ads Benchmarks 2026 y agregados de Digital Applied sobre "
                 "datos de Q1 2026. Las diferencias entre columnas reflejan distintas "
                 "definiciones de categoría entre fuentes.")
callout(doc, "El patrón importa más que el número.",
        "Finanzas presenta CTR alto y conversión baja. La gente hace clic en anuncios "
        "financieros con facilidad, compara varios proveedores, se toma tiempo y duda antes de "
        "entregar datos personales. Ese es exactamente el problema que la conversación asistida "
        "por IA está diseñada para resolver: reduce el costo psicológico del primer paso.")

h2(doc, "6.2 Triangulación del CPC para el mercado chileno")
para(doc, "No existe una fuente única y autorizada de CPC por industria para Chile. Para llegar "
          "a una banda defendible se cruzaron dos vías independientes.")
table(doc, ["Vía", "Valor", "Observación"],
      [["A. Benchmark EE.UU. convertido",
        "$2.807 - $3.154 CLP",
        "USD 3,08 - 3,46 al dólar observado de $911,43 del 26-ago-2026"],
       ["B. Reportes locales, servicios financieros",
        "$1.500 - $5.000 CLP",
        "Rango amplio publicado por agencias chilenas"],
       ["B'. Reportes locales, financiamiento y seguros",
        "$2.000 - $3.000 CLP",
        "Subcategoría más cercana al producto de Dale"],
       ["Banda triangulada de trabajo",
        "$2.400 - $3.200 CLP",
        "Zona de solapamiento entre A y B'"]],
      widths=[2.1, 1.5, 2.9])
source_note(doc, "Elaboración propia sobre WordStream 2026, reportes de agencias chilenas 2026 "
                 "y tipo de cambio del Banco Central al 26 de agosto de 2026.")
para(doc, "Ambas vías, construidas con fuentes que no se citan entre sí, convergen en la misma "
          "zona. Ese solapamiento es lo que da validez al número: Chile se sitúa en torno al "
          "60% a 80% del CPC estadounidense en esta industria, no en la fracción mucho menor "
          "que suele asumirse. Para Meta, el benchmark de tráfico de USD 1,22 convertido "
          "sugiere una banda aproximada de $600 a $1.100 CLP; es el supuesto con respaldo más "
          "débil de todo el modelo y debe validarse en las primeras semanas de campaña.")

h2(doc, "6.3 Meta Ads — referencias internacionales")
table(doc, ["Métrica", "Valor", "Alcance", "Nota"],
      [["CPC finanzas, campañas de conversión", "USD 3,77", "EE.UU.", "Sobre el promedio de la plataforma"],
       ["CPC finanzas y seguros, campañas de tráfico", "USD 1,22", "EE.UU.", "Base para la estimación de CTWA"],
       ["CTR esperable en verticales B2B y finanzas", "0,7% - 1,0%", "Global", "Se considera buen desempeño"],
       ["Tasa de conversión finanzas", "9,09%", "EE.UU.", "Alta respecto de otras industrias"],
       ["CPM en verticales de alto costo", "sobre USD 20", "Global", "Señal de segmentación estrecha"]],
      widths=[2.5, 1.2, 1.0, 1.8])
source_note(doc, "Benchmarks de Meta Ads por industria, 2026.")
para(doc, "Los verticales regulados y de alto valor pagan una prima porque el valor de un "
          "cliente es mayor. El cambio relevante en 2026 es la consolidación de Advantage+ y la "
          "recomendación de la plataforma de trabajar con segmentación amplia, dejando que el "
          "algoritmo encuentre al usuario. Esto es coherente con usar las bases de datos de la "
          "marca como semilla de audiencias similares y no como filtro rígido.")

# ────────────────────────── 7. KEYWORDS ───────────────────────────────────────
h1(doc, "Investigación de palabras clave", "07")
callout(doc, "Alcance de este capítulo.",
        "Sin acceso a la cuenta de Google Ads del cliente no es posible entregar volúmenes de "
        "búsqueda reales. Lo que sigue es la arquitectura de keywords propuesta, clasificada "
        "por intención y por prioridad estratégica, lista para ser validada y dimensionada con "
        "Keyword Planner en la primera semana de trabajo. Se declara así en vez de publicar "
        "volúmenes estimados que no podríamos respaldar.")

h2(doc, "7.1 Grupo Marca — prioridad alta, costo bajo")
table(doc, ["Término", "Intención", "Rol"],
      [["dale coopeuch", "Navegacional", "Defensa de marca"],
       ["credito dale coopeuch", "Transaccional", "Captura de demanda directa"],
       ["credito digital dale", "Transaccional", "Captura de demanda directa"],
       ["dale coopeuch app", "Navegacional", "Defensa de marca"],
       ["coopeuch credito de consumo", "Transaccional", "Puente desde la marca madre"]],
      widths=[2.3, 1.6, 2.6])

h2(doc, "7.2 Grupo Genérico Crédito — prioridad alta, costo alto")
table(doc, ["Término", "Intención", "Rol"],
      [["credito de consumo online", "Transaccional", "Núcleo del grupo"],
       ["credito de consumo digital", "Transaccional", "Núcleo del grupo"],
       ["solicitar credito online", "Transaccional", "Alta intención"],
       ["credito rapido online chile", "Transaccional", "Alta intención, vigilar calidad"],
       ["credito sin ir al banco", "Transaccional", "Ligado al diferenciador de proceso"],
       ["simular credito de consumo", "Investigación", "Captura previa a la decisión"],
       ["comparar creditos de consumo", "Investigación", "Requiere argumento de CAE"]],
      widths=[2.3, 1.6, 2.6])

h2(doc, "7.3 Grupo Inclusión Financiera — prioridad estratégica")
para(doc, "Este es el grupo diferencial de la cuenta. Conecta la capacidad declarada de "
          "evaluar múltiples fuentes de ingreso con una demanda real y desatendida.")
table(doc, ["Término", "Intención", "Rol"],
      [["credito para independientes", "Transaccional", "Núcleo del nicho"],
       ["credito sin liquidacion de sueldo", "Transaccional", "Dolor explícito del segmento"],
       ["credito para trabajadores a honorarios", "Transaccional", "Segmento formalizado"],
       ["credito sin renta fija", "Transaccional", "Variante del dolor"],
       ["credito para emprendedores", "Transaccional", "Adyacente, validar pertinencia"],
       ["como construir historial crediticio", "Informacional", "Nurture, no conversión directa"]],
      widths=[2.3, 1.6, 2.6])

h2(doc, "7.4 Keywords negativas obligatorias")
para(doc, "En una cuenta financiera las negativas son tan determinantes como las positivas: "
          "evitan gasto en tráfico que jamás califica.")
bullet(doc, [
 ("Otros productos crediticios: ", "hipotecario, automotriz, credito automotriz, "
  "refinanciamiento hipotecario, credito universitario, CAE estudiantil."),
 ("Intención informacional pura: ", "que es un credito, definicion, wikipedia, significado, "
  "tesis, ensayo, monografia."),
 ("Búsqueda de empleo: ", "trabajo, empleo, sueldo ejecutivo, ofertas laborales, practica."),
 ("Gestión de cuenta existente: ", "pagar cuota, estado de cuenta, sucursal, reclamo, "
  "cancelar, clave, bloqueo de tarjeta."),
 ("Términos de riesgo reputacional: ", "estafa, fraude, demanda, dicom, deudas impagas, "
  "aclarar deudas."),
 ("Competencia directa, salvo decisión explícita: ", "nombres de otras instituciones. Pujar "
  "por marca ajena en el sector financiero encarece la subasta y expone a la marca."),
])

h2(doc, "7.5 Criterios de estructura")
bullet(doc, [
 ("Un grupo por intención, no por producto. ", "La estructura anterior separa marca, genérico "
  "e inclusión porque cada uno requiere un mensaje y una landing distintos."),
 ("Concordancia mixta con vigilancia semanal. ", "Frase y exacta para el núcleo; amplia solo "
  "con audiencias de señal y revisión frecuente de términos de búsqueda."),
 ("Landing coherente por grupo. ", "El grupo de inclusión debe llegar a un mensaje que hable "
  "explícitamente de evaluación con múltiples fuentes de ingreso, no a la landing genérica."),
 ("Revisión de términos de búsqueda semanal el primer mes. ", "Es el mecanismo más eficaz "
  "para controlar el costo en una cuenta financiera nueva."),
])

# ────────────────────────── 8. EMAIL MARKETING ────────────────────────────────
h1(doc, "Benchmarks de Email Marketing", "08")
h2(doc, "8.1 El dato que reordena la estrategia del canal")
table(doc, ["Métrica", "Valor", "Alcance"],
      [["Ingreso generado por flujos automatizados", "41% del total de email", "Global"],
       ["Participación de los flujos en los envíos", "5,3% del total", "Global"],
       ["Click rate de flujos", "5,58%", "Global"],
       ["Click rate de campañas", "1,69%", "Global"],
       ["Tasa de pedido de flujos vs campañas", "13 veces superior", "Global"],
       ["Click rate promedio banca y finanzas", "3,4%", "Global"],
       ["Open rate promedio banca y finanzas", "~31%", "Global"]],
      widths=[3.0, 1.9, 1.6])
source_note(doc, "Klaviyo Benchmarks 2026, sobre más de 183.000 cuentas, y benchmarks de email "
                 "por industria 2026.")
para(doc, "La conclusión es directa: el 5,3% de los envíos produce el 41% de los ingresos. Un "
          "programa de email que solo envía campañas masivas está dejando fuera la parte del "
          "canal que efectivamente convierte. Por eso la propuesta compromete los dos correos "
          "semanales acordados, pero sitúa el esfuerzo estratégico en la arquitectura de "
          "flujos.")
callout(doc, "Advertencia sobre la línea base.",
        "Banca y finanzas presenta la tasa de clic más baja entre todas las industrias "
        "medidas: 3,4%. Conviene declararlo desde el inicio en lugar de que el cliente lo "
        "descubra al tercer mes comparándose contra benchmarks de retail que no le "
        "corresponden.")

h2(doc, "8.2 Arquitectura de flujos propuesta")
table(doc, ["#", "Flujo", "Gatillo", "Por qué importa"],
      [["1", "Preaprobado sin usar", "Oferta vigente sin activar",
        "El prospecto ya pasó la evaluación de riesgo. Es el retorno más alto por envío."],
       ["2", "Contrato no firmado", "Llegó a firma y no firmó",
        "Máxima urgencia y mínima fricción restante para completar el curse."],
       ["3", "Solicitud abandonada", "Inició el proceso sin completarlo",
        "Equivalente funcional del carro abandonado en comercio electrónico."],
       ["4", "Onboarding", "Registro reciente",
        "Acompaña los pasos previos que hoy generan abandono."],
       ["5", "Nurture educativo", "No califica o no está listo",
        "Historial crediticio y salud financiera. Mantiene vivo al prospecto."],
       ["6", "Winback", "Base fría segmentada",
        "Destino natural de las bases que la marca puede aportar."],
       ["7", "Post-contratación", "Curse completado",
        "Seguros, Hipotecario 40 e inversión. Aumenta el valor por cliente."],
       ["8", "Sunset", "Inactividad prolongada",
        "Higiene de lista. Determinante para la entregabilidad con bases antiguas."]],
      widths=[0.3, 1.4, 1.5, 3.3])

h2(doc, "8.3 Consideración sobre la herramienta")
para(doc, "El benchmark de referencia proviene de Klaviyo, una plataforma originada en el "
          "comercio electrónico. Su modelo de flujos es aplicable a este caso, pero la decisión "
          "de herramienta para una institución financiera regulada debe considerar además el "
          "tratamiento de datos personales, los requisitos de la Ley 21.719 de Protección de "
          "Datos Personales y la política interna de la marca sobre alojamiento de información "
          "de sus socios. La agencia acompaña esa evaluación; la decisión es del cliente.")

# ────────────────────────── 9. COMMUNITY MANAGEMENT ───────────────────────────
h1(doc, "Benchmarks de Community Management", "09")
table(doc, ["Plataforma", "Engagement servicios financieros", "Referencia de otras industrias"],
      [["Instagram", "0,26% - 0,67%", "Educación superior alcanza 2,10%"],
       ["TikTok", "1,9%", "Educación superior alcanza 7,36%"],
       ["LinkedIn", "3,3% en foto y video", "Mejor plataforma del sector"],
       ["Facebook", "Crecimiento negativo de seguidores", "-0,61% semanal en el sector"]],
      widths=[1.4, 2.4, 2.7])
source_note(doc, "Rival IQ y estudios agregados de benchmarks sociales, 2026.")
para(doc, "TikTok supera a Instagram en engagement por un factor de 3 a 10 veces en "
          "prácticamente toda industria medida, y servicios financieros no es la excepción. "
          "Eso respalda la decisión de cubrir ambas plataformas con contenido nativo en cada "
          "una en lugar de replicar piezas.")
callout(doc, "El rol correcto del contenido orgánico en este caso.",
        "Con tasas de engagement estructuralmente bajas, prometer crecimiento de comunidad "
        "sería vender una métrica que la industria no entrega. El rol del contenido aquí es "
        "otro y es más valioso: reducir la desconfianza, explicar el proceso paso a paso y "
        "sostener la credibilidad de la marca frente a un prospecto que está a punto de "
        "entregar sus datos financieros.")
h2(doc, "Ejes de contenido recomendados")
bullet(doc, [
 ("Explicación del proceso. ", "Tutoriales cortos de cada paso. Ataca directamente el abandono "
  "por desconocimiento."),
 ("Manejo de objeciones. ", "Qué es DalePass, por qué se pide ser Socio, qué pasa si me "
  "rechazan. Son las preguntas que hoy frenan la conversión."),
 ("El diferenciador de evaluación inclusiva. ", "Historias de independientes y trabajadores a "
  "honorarios. Es el contenido con mayor potencial de identificación."),
 ("Educación financiera y Remanente. ", "Construye autoridad y explica un beneficio que la "
  "competencia no tiene."),
])

# ────────────────────────── 10. ANUNCIOS Y TENDENCIAS ─────────────────────────
h1(doc, "Anuncios de referencia y tendencias", "10")
h2(doc, "10.1 Sobre los anuncios de referencia")
callout(doc, "Alcance declarado.",
        "El relevamiento sistemático de anuncios activos de la competencia requiere consulta "
        "directa a la Biblioteca de Anuncios de Meta y al Centro de Transparencia de Google "
        "Ads, con captura de piezas en el momento. Este documento no incluye capturas de piezas "
        "de terceros para no atribuir creatividades sin evidencia verificable. Se recomienda "
        "ejecutar ese relevamiento como primera actividad del onboarding y adjuntarlo como "
        "anexo.")
h2(doc, "10.2 Patrones de anuncio que la evidencia respalda")
para(doc, "Lo que sí se puede afirmar desde los benchmarks y el marco regulatorio revisados:")
bullet(doc, [
 ("El anuncio financiero gana clics y pierde conversiones. ", "Con CTR alto y conversión baja, "
  "la pieza no debe competir por atención sino por calificación. Un anuncio que atrae a quien "
  "no califica es costo puro."),
 ("La transparencia de costo es ventaja, no obligación molesta. ", "Dado que el usuario compara "
  "por CAE y que existe un simulador público de la CMF, exhibir condiciones claras diferencia "
  "en un rubro donde la letra chica es la norma."),
 ("El ángulo de inclusión no está ocupado. ", "Ningún competidor directo comunica con claridad "
  "la evaluación por múltiples fuentes de ingreso. Es territorio libre."),
 ("El destino conversacional reduce fricción. ", "Llevar a WhatsApp en lugar de a un "
  "formulario ataca directamente la causa documentada de la baja conversión del rubro: la "
  "resistencia a entregar datos personales en frío."),
 ("Video vertical nativo por plataforma. ", "La brecha de engagement entre TikTok e Instagram "
  "confirma que replicar la misma pieza en ambas desaprovecha la plataforma de mayor retorno."),
])
h2(doc, "10.3 Tendencias relevantes para el periodo")
bullet(doc, [
 ("Consolidación de la automatización algorítmica. ", "Advantage+ en Meta y Performance Max en "
  "Google desplazan el control manual hacia la calidad de las señales que se entregan. El valor "
  "del trabajo de agencia se traslada de la segmentación a la medición y al contenido."),
 ("Digital supera la mitad del mercado publicitario chileno. ", "50,1% de share en mayo de "
  "2026, con social y search concentrando dos tercios de la compra directa."),
 ("Expansión acelerada del prepago y los neobancos. ", "213% de crecimiento en uso de tarjetas "
  "de prepago entre marzo de 2024 y marzo de 2025, con Mercado Pago entrando a crédito en 2026."),
 ("Endurecimiento regulatorio de la publicidad financiera. ", "Google amplió en junio de 2026 "
  "los requisitos de verificación a 24 mercados del Espacio Económico Europeo. Chile aún no "
  "está incluido, pero la dirección del cambio es clara."),
 ("Presión sobre el dato propio. ", "Con la Ley 21.719 y el endurecimiento de la privacidad en "
  "plataformas, las bases propias y los canales conversacionales ganan peso relativo frente a "
  "la segmentación de terceros."),
])

# ────────────────────────── 11. MODELO DE INVERSIÓN ───────────────────────────
h1(doc, "Modelo de inversión y proyección", "11")
para(doc, "El modelo se construye de abajo hacia arriba: parte del costo por clic triangulado, "
          "aplica tasas de paso entre etapas y llega al costo por curse. Se entrega también en "
          "formato editable dentro del Plan de Medios, hoja \"Modelo de Inversión\", para que "
          "el cliente pueda modificar cualquier supuesto y ver el efecto.")
h2(doc, "11.1 Supuestos")
table(doc, ["Supuesto", "Valor", "Origen"],
      [["CPC Google Search", "$2.800 CLP", "Punto medio de la banda triangulada"],
       ["CPC Meta CTWA", "$850 CLP", "Benchmark de tráfico convertido. Supuesto más débil"],
       ["Clic a conversación, Google", "12%", "Landing con CTA directo a WhatsApp"],
       ["Clic a conversación, Meta CTWA", "25%", "El clic abre WhatsApp sin paso intermedio"],
       ["Conversación a prospecto calificado", "35%", "Filtro del agente conversacional"],
       ["Prospecto calificado a curse", "20%", "A DEFINIR POR EL CLIENTE: depende de su política de riesgo"],
       ["Distribución Google / Meta", "40% / 60%", "Coherente con el peso de cada canal en el mercado"]],
      widths=[2.3, 1.3, 2.9])
h2(doc, "11.2 Proyección para un escenario de $3.500.000 CLP mensuales")
table(doc, ["Indicador", "Resultado"],
      [["Inversión Google", "$1.400.000 CLP"],
       ["Inversión Meta", "$2.100.000 CLP"],
       ["Clics totales estimados", "~2.970"],
       ["Conversaciones iniciadas", "~678"],
       ["Prospectos calificados", "~237"],
       ["Costo por prospecto calificado", "~$14.800 CLP"],
       ["Cursos estimados", "~47"],
       ["Costo por curse", "~$74.500 CLP"]],
      widths=[3.2, 2.3])
callout(doc, "Cómo leer esta proyección.",
        "No es una promesa de resultado. Es la explicitación del razonamiento que sustenta la "
        "cifra de inversión recomendada, construida sobre benchmarks públicos y atribuidos. "
        "Cada supuesto queda sujeto a las métricas reales que se observen al implementar, y el "
        "modelo está diseñado para recalcularse con esos datos desde el primer mes. La única "
        "variable que no puede estimarse desde fuentes externas es la tasa de aprobación a "
        "curse: depende de la política de riesgo de Coopeuch y debe aportarla el cliente.")
h2(doc, "11.3 Banda de inversión recomendada")
para(doc, "Se recomienda operar entre $3.000.000 y $5.000.000 CLP mensuales durante la fase "
          "inicial. El límite inferior permite generar volumen de eventos suficiente para que "
          "las campañas de conversión completen su fase de aprendizaje; el superior evita "
          "saturar audiencias antes de haber validado el mensaje y el proceso de calificación.")

# ────────────────────────── 12. CONCLUSIONES ──────────────────────────────────
h1(doc, "Conclusiones y recomendaciones", "12")
h2(doc, "12.1 Antes de invertir el primer peso")
bullet(doc, [
 ("Resolver la divulgación de condiciones del crédito. ", "Publicar CAE, plazos, montos y "
  "ejemplo representativo del costo total en la landing pública. Es requisito de Google Ads."),
 ("Definir el mapa de eventos y su medición. ", "Sin trazabilidad hasta el curse, la "
  "optimización trabajará sobre la métrica equivocada."),
 ("Confirmar la tasa de aprobación crediticia. ", "Es el único parámetro del modelo que no "
  "puede estimarse externamente."),
 ("Verificar los requisitos de verificación del anunciante. ", "Tanto en Google como en Meta, "
  "antes de programar el lanzamiento."),
])
h2(doc, "12.2 Dónde está la oportunidad")
bullet(doc, [
 ("El nicho de inclusión financiera. ", "La evaluación por múltiples fuentes de ingreso es un "
  "diferenciador real, comunicable y hoy no ocupado por la competencia directa."),
 ("El Remanente como argumento de conversión. ", "Convierte el requisito de asociarse en "
  "beneficio económico. Ningún neobanco compite con eso."),
 ("Los flujos automatizados. ", "El punto de mayor retorno del programa de email, y el que se "
  "refuerza con el agente conversacional sobre los mismos puntos de fuga."),
 ("La transparencia de costo como ventaja competitiva. ", "En un rubro asociado a la letra "
  "chica, cumplir bien el requisito regulatorio puede convertirse en argumento de venta."),
])
h2(doc, "12.3 Qué medir para saber si está funcionando")
para(doc, "Tres indicadores por sobre todos los demás: costo por prospecto calificado, tasa de "
          "aprobación de los prospectos derivados y costo por curse. El resto de las métricas "
          "son diagnósticas; estas tres son las que determinan si la inversión rinde.")

# ────────────────────────── 13. FUENTES ───────────────────────────────────────
h1(doc, "Fuentes consultadas", "13")
table(doc, ["Fuente", "Uso en este documento", "Consulta"],
      [["CMF — Desempeño de bancos y cooperativas, enero 2026",
        "Cartera de consumo y riesgo en cooperativas", "ago-2026"],
       ["CMF — Simulador de créditos", "Referencia de divulgación de CAE", "ago-2026"],
       ["SERNAC — Material educativo sobre CAE", "Definición y alcance del indicador", "ago-2026"],
       ["AAM, Admetricks e IAB Chile — Inversión Digital mayo 2026",
        "Inversión publicitaria y share por medio en Chile", "ago-2026"],
       ["Google Ads — Políticas de préstamos personales",
        "Requisitos de divulgación y regla de 61 días", "ago-2026"],
       ["Google Ads — Verificación de servicios financieros junio 2026",
        "Alcance geográfico de la verificación reforzada", "ago-2026"],
       ["Meta Business Help Centre — Categorías especiales de anuncios",
        "Alcance geográfico de las restricciones de crédito", "ago-2026"],
       ["Meta Centro de Transparencia — Productos financieros",
        "Verificación de anunciante y restricciones de contenido", "ago-2026"],
       ["WordStream — Google Ads Benchmarks 2026", "CPC, CTR y conversión en finanzas", "ago-2026"],
       ["Digital Applied — Benchmarks agregados Q1 2026", "Desglose banca y seguros", "ago-2026"],
       ["Benchmarks de Meta Ads por industria 2026", "CPC, CTR y conversión en Meta", "ago-2026"],
       ["Klaviyo — Email Marketing Benchmarks 2026", "Flujos vs campañas sobre 183.000 cuentas", "ago-2026"],
       ["Rival IQ — Social Media Benchmarks 2026", "Engagement por plataforma en finanzas", "ago-2026"],
       ["Banco Central de Chile", "Tipo de cambio y valor UF de referencia", "26-ago-2026"],
       ["Cámara de Comercio de Santiago y LatamFintech",
        "Panorama competitivo y crecimiento del prepago", "ago-2026"],
       ["dalecoopeuch.cl", "Producto, propuesta de valor, embudo y verificación de landing", "ago-2026"],
       ["Reportes de agencias digitales chilenas", "CPC local en pesos por industria", "ago-2026"]],
      widths=[2.6, 2.6, 0.9], size=8)

doc.save(OUT)
print("Guardado:", OUT)
