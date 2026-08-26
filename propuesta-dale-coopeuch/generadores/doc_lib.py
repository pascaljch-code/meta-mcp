# -*- coding: utf-8 -*-
"""Helpers de estilo Intothecom para python-docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ORANGE = RGBColor(0xE5, 0x70, 0x00)
INK    = RGBColor(0x26, 0x26, 0x26)
GRAY   = RGBColor(0x59, 0x59, 0x59)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
FONT   = "Be Vietnam Pro"


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)


def setup(doc):
    st = doc.styles['Normal']
    st.font.name = FONT
    st.font.size = Pt(10)
    st.font.color.rgb = INK
    st.element.rPr.rFonts.set(qn('w:eastAsia'), FONT)
    pf = st.paragraph_format
    pf.space_after = Pt(6); pf.line_spacing = 1.25
    for s in doc.sections:
        s.top_margin = s.bottom_margin = Cm(2.2)
        s.left_margin = s.right_margin = Cm(2.4)
    return doc


def h1(doc, text, num=None):
    doc.add_page_break()
    if num:
        p = doc.add_paragraph()
        r = p.add_run(num)
        r.font.name = FONT; r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = ORANGE
        p.paragraph_format.space_after = Pt(0)
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(19); r.font.bold = True; r.font.color.rgb = INK
    p.paragraph_format.space_after = Pt(10)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = FONT; r.font.size = Pt(12.5); r.font.bold = True
    r.font.color.rgb = ORANGE
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    return p


def para(doc, text, size=10, color=INK, bold=False, italic=False, after=6, align=None):
    p = doc.add_paragraph()
    if isinstance(text, str):
        text = [(text, {})]
    for t, ov in text:
        r = p.add_run(t)
        r.font.name = FONT
        r.font.size = Pt(ov.get("size", size))
        r.font.bold = ov.get("bold", bold)
        r.font.italic = ov.get("italic", italic)
        r.font.color.rgb = ov.get("color", color)
    p.paragraph_format.space_after = Pt(after)
    if align:
        p.alignment = align
    return p


def bullet(doc, items, size=10):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(it, tuple):
            head, rest = it
            r = p.add_run(head)
            r.font.name = FONT; r.font.size = Pt(size); r.font.bold = True
            r.font.color.rgb = INK
            r2 = p.add_run(rest)
            r2.font.name = FONT; r2.font.size = Pt(size); r2.font.color.rgb = GRAY
        else:
            r = p.add_run(it)
            r.font.name = FONT; r.font.size = Pt(size); r.font.color.rgb = GRAY
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2


def table(doc, headers, rows, widths=None, size=8.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = FONT; r.font.size = Pt(size); r.font.bold = True
        r.font.color.rgb = WHITE
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        shade(hdr[i], "3A3838")
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(v))
            r.font.name = FONT; r.font.size = Pt(size)
            r.font.color.rgb = INK
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            if ri % 2 == 1:
                shade(cells[i], "F7F7F7")
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def callout(doc, title, text):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    shade(c, "FDF0E4")
    c.text = ""
    p = c.paragraphs[0]
    r = p.add_run(title + "  ")
    r.font.name = FONT; r.font.size = Pt(9.5); r.font.bold = True; r.font.color.rgb = ORANGE
    r2 = p.add_run(text)
    r2.font.name = FONT; r2.font.size = Pt(9.5); r2.font.color.rgb = INK
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def source_note(doc, text):
    para(doc, [("Fuente: ", {"bold": True, "size": 8, "color": GRAY}),
               (text, {"size": 8, "color": GRAY, "italic": True})], after=10)
